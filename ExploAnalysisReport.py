from math import e
from pydoc import doc
from turtle import width
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from numpy import add, median
from requests import get
from fonction import *
from FonctionRapport import *
import pandas as pd
import os
import time
from scipy.stats import skew, kurtosis  
import warnings

df = pd.read_pickle('ProcessedDonnéesDelliaux.pkl')
df.rename(columns={'V1_Complications_Post_Operatoires_Reanimation_Rapport_PaO2/FiO2': 'V1_Complications_Post_Operatoires_Reanimation_Rapport_PaO2_FiO2', "V0_Maladie_Initiale_Nan": "V0_Maladie_Initiale_Aucune", "V1_Protocole_Immunosuppression_Nan": "V1_Protocole_Immunosuppression_Aucune", 'V1_Donneur_GDS_Rapport_Pa02/Fi02': 'V1_Donneur_GDS_Rapport_Pa02_Fi02', 'V2VX_Complications_Infectieuses_"Evenement_Infectieux_Dorigine_Bacterienne"':'V2VX_Complications_Infectieuses_Evenement_Infectieux_Dorigine_Bacterienne'}, inplace = True)
df.replace([np.inf, -np.inf], np.nan, inplace=True)
# df_unique = create_unique_df(df)
df_unique = pd.read_pickle('df_patient.pkl')
df_unique.rename(columns={'V1_Complications_Post_Operatoires_Reanimation_Rapport_PaO2/FiO2': 'V1_Complications_Post_Operatoires_Reanimation_Rapport_PaO2_FiO2', "V0_Maladie_Initiale_Nan": "V0_Maladie_Initiale_Aucune", "V1_Protocole_Immunosuppression_Nan": "V1_Protocole_Immunosuppression_Aucune", 'V1_Donneur_GDS_Rapport_Pa02/Fi02': 'V1_Donneur_GDS_Rapport_Pa02_Fi02', 'V2VX_Complications_Infectieuses_"Evenement_Infectieux_Dorigine_Bacterienne"':'V2VX_Complications_Infectieuses_Evenement_Infectieux_Dorigine_Bacterienne'}, inplace = True)

# print(df_unique.info())

start_time = time.time()
doc = Document()


# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
heading = doc.add_heading(level=0)
run = heading.add_run('-ANALYSE DES DONNEES COLT-')
run.bold = True
doc.add_paragraph()
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_heading('I - Structure de la base de données', level=0)
# print(f"forme de la base de données est {df.shape}")
doc.add_heading('Aperçu de la base de données', level=1)
doc.add_paragraph(f"Nbre lignes: {df.shape[0]}"
                 f"\nNbre de colonnes: {df.shape[1]}"
                 f"\nNbre de cellules: {df.shape[0]*df.shape[1]}" 
                 f"\nNbre de cellules vides: {df.isna().sum().sum()}"
                 f"\nPourcentage de cellules vides: {round(df.isna().sum().sum()*100/(df.shape[0]*df.shape[1]), 2)} %"
                 f"\nNbre de patients: {df['Admin_Identifiant_Colt'].nunique()}"
                 )

doc.add_heading("Type de données", level = 1)
doc.add_paragraph(f"Alphanumériques: {df.select_dtypes(include=['object']).shape[1]}  "
                  f"\nBooléennes: {df.select_dtypes(include=['boolean']).shape[1]}"
                  f"\nCatégorielles: {df.select_dtypes(include=['category']).shape[1]}  "
                  f"\nDates: {df.select_dtypes(include=['datetime64']).shape[1]} "
                  f"\nNumériques: {df.select_dtypes(include=['float64', 'int64']).shape[1]} "  
                  f"\nNbre Total: {df.select_dtypes(include = ['object']).shape[1] + df.select_dtypes(include=['category']).shape[1] + df.select_dtypes(include=['boolean']).shape[1]
                                    + df.select_dtypes(include=['datetime64']).shape[1] + df.select_dtypes(include=['float64', 'int64']).shape[1]}"                
                  )

doc.add_heading("Structure", level=1)
doc.add_picture('tree.png', width=Inches(3.7))

# df = df.sort_index(axis=1)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('II - Descriptif Univarié', level=0)
doc.add_heading('Admin', level=1)
display_v(df_unique, doc, 'Admin')
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# doc.add_heading()
prefixes_2 = ["Antecedant", "Colonisation_Bact", "Colonisation_Champignon", "Colonisation_MycoBact", "EFR", "GDS", "Generales", "Histocompatibilite", "Maladie_Initiale", "Serologies", "Traitement", "Ventillation"]
# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('V0', level=1)
# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
for prefix in prefixes_2:
    chosen_df = df if prefix == 'Fonction_Respiratoire' else df_unique
    doc.add_heading(prefix.replace("_"," "), level=1)
    doc.add_paragraph()
    display_v(chosen_df, doc, 'V0', prefix)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Synthèse Admin', level=1)
doc.add_picture('V0_Antecedant_Cancer_Si_Oui.png', width=Inches(6))
doc.add_picture('V0_Antecedant_Cancer_Si_Oui.png', width=Inches(6))

prefixes_v1_1 = ["Complications", "Donneur", "Procedure_chir", "Protocole_Immunosuppression"]
prefixes_v1_2 = ["Per_Operatoires", "Post_Operatoires", "GDS", "Histocompatibilite", "Procedure", "Serologies"]
prefixes_v1_3 = ["Crossmatch", "Prelevement", "Precoces", "Reanimation"]

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('V1', level=1)
# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
for prefix1 in prefixes_v1_1:
    if prefix1 == "Complications":
        lst = ["Per_Operatoires", "Post_Operatoires"]
        for prefix2 in lst:   
        # doc.add_heading(prefix1.replace("_", " "), level=2)
            if prefix2 == "Post_Operatoires":
                lst1 = ["Precoces", "Prelevement", "Reanimation"]
                for prefix3 in lst1:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(f"{prefix1.replace('_',' ')} {prefix2.replace('_',' ')} {prefix3}")
                    run.bold = True
                    paragraph.style = doc.styles['Heading 2']
                    doc.add_paragraph()
                    display_v(df_unique, doc, 'V1', prefix1, prefix2, prefix3)
                    doc.add_paragraph()
                    # doc.add_paragraph()
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{prefix1.replace('_',' ')} {prefix2.replace('_',' ')}")
                run.bold = True
                paragraph.style = doc.styles['Heading 2']
                doc.add_paragraph()
                display_v(df_unique, doc, 'V1', prefix1, prefix2)

                doc.add_paragraph()
                # doc.add_paragraph()

    elif prefix1 == "Donneur":
        lst = ["GDS", "Histocompatibilite", "Procedure", "Serologies"]
        # doc.add_heading(prefix1, level=2)
        for prefix2 in lst:
            if prefix2 == "Procedure":
                lst1 = ["Crossmatch", "Prelevement"]
                for prefix3 in lst1:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(f"{prefix1} {prefix2} {prefix3}")
                    run.bold = True
                    paragraph.style = doc.styles['Heading 2']
                    display_v(df_unique, doc, 'V1', prefix1, prefix2, prefix3)
                    doc.add_paragraph()
                    # doc.add_paragraph()
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{prefix1.replace('_',' ')} {prefix2.replace('_',' ')}")
                run.bold = True
                paragraph.style = doc.styles['Heading 2']
                display_v(df_unique, doc, 'V1', prefix1, prefix2)
                doc.add_paragraph()
                # doc.add_paragraph()
            
    elif prefix1 == "Procedure_chir":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(prefix1.replace("_", " "))
        run.bold = True
        paragraph.style = doc.styles['Heading 2']
        display_v(df_unique, doc, 'V1', prefix1)
        doc.add_paragraph()
        # doc.add_paragraph()
    elif prefix1 == "Protocole_Immunosuppression":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(prefix1.replace("_", " "))
        run.bold = True 
        paragraph.style = doc.styles['Heading 2']
        display_v(df_unique, doc, 'V1', prefix1)
        doc.add_paragraph()
        # doc.add_paragraph()
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(prefix1.replace("_ ", " "))
        run.bold = True 
        paragraph.style = doc.styles['Heading 2']
        display_v(df, 'V1', prefix1, None, None, doc)
        doc.add_paragraph()
        # doc.add_paragraph()

doc.add_heading('V2VX', level=1)
prefix_v2_1 = ["Rejet", "Complications", "Suivi_Traitement",'Bio' "Fonction_Respiratoire", "Survie"]


display_v(df_unique, doc, 'V2VX', 'Centre')
doc.add_paragraph()
# doc.add_paragraph()

for prefix in prefix_v2_1:
    # doc.add_heading(prefix.replace("_"," "), level=2)
    if prefix == "Rejet":
        doc.add_heading(prefix.replace("_"," "), level=2)
        display_v(df_unique, doc, 'V2VX', prefix)
        doc.add_paragraph()
        # doc.add_paragraph()
        
    elif prefix == "Complications":
        # doc.add_heading(prefix, level=3)
        lst2 = ["Bronchiques", "Vasculaires", "Infectieuses", "Non_Infectieuses"]
        for prefix2 in lst2:
            doc.add_heading(f"{prefix.replace("_"," ")} {prefix2.replace("_", " ")}", level=2)
            display_v(df_unique, doc, 'V2VX', prefix, prefix2)
            doc.add_paragraph()
            # doc.add_paragraph()
            
    elif prefix == "Suivi_Traitement":
        doc.add_heading(prefix.replace("_"," "), level=2)
        display_v(df_unique, doc, 'V2VX', prefix)
        doc.add_paragraph()
        # doc.add_paragraph()
        

    elif prefix == "Fonction_Respiratoire":
        doc.add_heading(prefix.replace("_"," "), level=2)
        display_v(df, doc, 'V2VX', prefix)
        doc.add_paragraph()
        # doc.add_paragraph()
        

    elif prefix == "Survie":
        doc.add_heading(prefix.replace("_"," "), level=2)
        display_v(df_unique, doc, 'V2VX', prefix)
# print('Traitement des variable terminé....')

paragraph = doc.add_paragraph()
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# colors = sns.color_palette("hsv", 10)

# doc.add_picture(plot_pie(df_unique, 'V1_Procedure_chir_Type_de_Greffe', colors))
# paragraph.add_run().add_picture('Type_de_Greffe.png',  width=Inches(4))
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('ANNEXES', level=0)
doc.add_heading('Annexe 1.A - Admin', level=1)
add_summary_to_doc(df_unique, 'Admin', doc)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Annexe 1.0 - V0', level=1)
for prefix in prefixes_2:
    doc.add_heading(prefix.replace("_"," "), level=2)
    add_summary_to_doc(df_unique, 'V0', doc, prefix)
    doc.add_paragraph()
    # doc.add_paragraph()


doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Annexe 1.1 - V1', level=1)
for prefix1 in prefixes_v1_1:
    if prefix1 == "Complications":
        lst = ["Per_Operatoires", "Post_Operatoires"]
        for prefix2 in lst:   
        # doc.add_heading(prefix1.replace("_", " "), level=2)
            if prefix2 == "Post_Operatoires":
                lst1 = ["Precoces", "Prelevement", "Reanimation"]
                for prefix3 in lst1:
                    doc.add_heading(f"{prefix1.replace("_"," ")} {prefix2.replace("_"," ")} {prefix3}", level=2)
                    add_summary_to_doc(df_unique, 'V1',doc, prefix1, prefix2, prefix3)
                    doc.add_paragraph()
                    # doc.add_paragraph()
            else:
                doc.add_heading(f"{prefix1.replace("_"," ")} {prefix2.replace("_"," ")}", level=2)
                add_summary_to_doc(df, 'V1', doc, prefix1, prefix2)
                doc.add_paragraph()
                # doc.add_paragraph()

    elif prefix1 == "Donneur":
        lst = ["GDS", "Histocompatibilite", "Procedure", "Serologies"]
        # doc.add_heading(prefix1, level=2)
        for prefix2 in lst:
            if prefix2 == "Procedure":
                lst1 = ["Crossmatch", "Prelevement"]
                for prefix3 in lst1:
                    doc.add_heading(f"{prefix1} {prefix2} {prefix3}", level=2)
                    add_summary_to_doc(df_unique, 'V1', doc, prefix1, prefix2, prefix3)
                    doc.add_paragraph()
                    # doc.add_paragraph()
            else:
                doc.add_heading(f"{prefix1.replace("_", " ")} {prefix2.replace("_", " ")}", level=2)
                add_summary_to_doc(df_unique, 'V1', doc, prefix1, prefix2, None)
                doc.add_paragraph()
                # doc.add_paragraph()
            
    elif prefix1 == "Procedure_chir":
        doc.add_heading(prefix1.replace("_", " "), level=2)
        add_summary_to_doc(df_unique, 'V1',doc, prefix1, None, None)
        doc.add_paragraph()
        # doc.add_paragraph()
    elif prefix1 == "Protocole_Immunosuppression":
        doc.add_heading(prefix1.replace("_", " "), level=2)
        add_summary_to_doc(df_unique, 'V1',doc, prefix1, None, None)
        doc.add_paragraph()
        # doc.add_paragraph()
    else:
        doc.add_heading(prefix1.replace("_ ", " "), level=2)
        add_summary_to_doc(df_unique, 'V1', doc, prefix1, None, None)
        doc.add_paragraph()
        # doc.add_paragraph()
    
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Annexe 1.2 - V2VX', level=1)
prefix_v2_1 = ["Rejet", "Complications", "Suivi_Traitement", "Fonction_Respiratoire", "Survie"]

add_summary_to_doc(df, 'V2VX', doc, 'Centre', None, None)
doc.add_paragraph()
# doc.add_paragraph()

for prefix in prefix_v2_1:
    # doc.add_heading(prefix.replace("_"," "), level=2)
    if prefix == "Rejet":
        doc.add_heading(prefix.replace("_"," "), level=2)
        add_summary_to_doc(df_unique, 'V2VX', doc, prefix, None, None)
        doc.add_paragraph()
        # doc.add_paragraph()
        
    elif prefix == "Complications":
        # doc.add_heading(prefix, level=3)
        lst2 = ["Bronchiques", "Vasculaires", "Infectieuses", "Non_Infectieuses"]
        for prefix2 in lst2:
            doc.add_heading(f"{prefix.replace("_"," ")} {prefix2.replace("_", " ")}", level=2)
            add_summary_to_doc(df_unique, 'V2VX', doc, prefix, prefix2, None)
            doc.add_paragraph()
            # doc.add_paragraph()
            
    elif prefix == "Suivi_Traitement":
        doc.add_heading(prefix.replace("_"," "), level=2)
        add_summary_to_doc(df_unique, 'V2VX', doc, prefix, None, None)
        doc.add_paragraph()
        # doc.add_paragraph()
        

    elif prefix == "Fonction_Respiratoire":
        doc.add_heading(prefix.replace("_"," "), level=2)
        add_summary_to_doc(df, 'V2VX', doc, prefix, None, None)
        doc.add_paragraph()
        # doc.add_paragraph()
        

    elif prefix == "Survie":
        doc.add_heading(prefix.replace("_"," "), level=2)
        add_summary_to_doc(df_unique, 'V2VX', doc, prefix, None, None)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Annexe 1.B - Détails des methods utilisées', level=1)
doc.add_heading(' Teste de normalité: ', level=2)
# Insérer le contenu de l'explication
explication_bootstrap = """Pour evaluer la normalité, un échantillon aléatoire de 20 à 50 données est prélevé sur l'ensemble des données observées.\
Ensuite on applique 3 (trois) tests : le test de Shapiro-Wilk, le test de D'Agostino et Pearson, ainsi que le test d'Anderson-Darling a cette échantillon.\
Puis on catégorise le resulat comme 'Gaussien" si deux des trois tests indiquent que l'échantillon est "Gaussien" si non le resulat est categorisé comme "Non Gaussien". 
Ce procedé est répété mille fois et résultat final de l'évaluation de la normalité est déterminé en comptant le nombre d'itérations où chaque état ("Gaussien" ou "Non Gaussien") apparaît.\
Si la majorité des itérations indiquent que l'échantillon est "Gaussien", il est considéré comme tel, sinon, il est considéré comme "Non Gaussien".
"""

paragraph = doc.add_paragraph(explication_bootstrap)
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# paragraph.runs

doc.add_heading(' Detection des outliers: ', level=2)

detailed_explanation = """Pour la détection des outliers, nous avons utilisé quatre (4) méthodes différentes : Local Outlier Factor (LOF), Isolation Forest \
et Elliptic Envelope et celle de la methode du boxplot. Nous avons collécté les outliers detectés par chaque methode. Puis on a identifié les outliers communs \
que nous avons mis en rouge sur le boxplot.
"""
paragraph = doc.add_paragraph(detailed_explanation)
# paragraph.runs
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# liste des variables dans l'ordre qu'elles apparaisse dans la base de données
# en respectant le prefixe

colonnes_dans_l_ordre = df.columns.tolist()
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Annexe 1.C - Liste des variables dans l\'ordre de la base de données', level=1)
doc.add_paragraph()

create_liste_colnames(df, doc)

# doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)# Dejà sur une nouvelle page 
# page break
doc.add_page_break()
doc.add_heading('Annexe 1.D - Liste des variables crées à partir des autres', level=1)
table = doc.add_table(rows=6, cols=5)
# table.style = 'Table Grid'
table.cell(0,0).merge(table.cell(0,1)).merge(table.cell(0,2)).merge(table.cell(0,3))
cell = table.cell(0,0)
para = cell.paragraphs[0]
run = para.add_run('Prefixes')
run.bold = True
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell = table.cell(0,4)
para = cell.paragraphs[0]
run = para.add_run('Variables')
run.bold = True
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table.cell(1,0).text = 'V0'
table.cell(1,1).text = 'Maladie_Initiale'
table.cell(1,4).text = 'Aucune'

table.cell(2,0).text = 'Admin'
table.cell(2,4).text = 'Age Greffe'
table.cell(3,0).text = 'V1'
table.cell(3,4).text = 'Delai Greffe Deces'
table.cell(4,0).text = 'V2VX'
table.cell(4,4).text = 'Age Deces'

table.cell(5,0).text = 'V2VX'
table.cell(5,2).text = 'Fonction Respiratoire'
table.cell(5,4).text = 'Serie Chronologique'


doc.save('ExploAnalysisReport.docx')
os.startfile('ExploAnalysisReport.docx')

print(f"--- {(time.time() - start_time)/60} minutes ---")