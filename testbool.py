from tracemalloc import start
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from matplotlib import colors
import seaborn as sns
import numpy as np
import os, time, pandas as pd, matplotlib.pyplot as plt
from scipy.stats import chi2_contingency, chisquare
from docx.enum.section import WD_ORIENT
from statsmodels.stats.multitest import multipletests
start_time = time.time()
# df = pd.read_pickle(r"ProcessedDonnéesDelliaux.pkl")
df_unique = pd.read_pickle(r"df_patient.pkl")

doc = Document()
save_dir = 'Plots of distribution'
import numpy as np
from scipy.stats import chi2_contingency


def chi2_summary(data, variable, centre_suivi_categories):
    try:
        # Create the contingency table
        contingency_table = []
        unique_values = data[variable].dropna().unique()
        for centre in centre_suivi_categories:
            data_category = data[data['V2VX_Centre_Suivi'] == centre][variable].dropna()
            observed_frequencies = data_category.value_counts().reindex(unique_values, fill_value=0).sort_index()
            contingency_table.append(observed_frequencies)
        
        # Convert the list to a numpy array
        try:
            contingency_table = np.array(contingency_table)
        
        # Perform the Chi-squared test
            chi2_stat, p_val_chi2, dof, expected = chi2_contingency(contingency_table)
        except Exception as e:
            print(f"Error: {e}")
            return f"Error: {e}"
        
        # Check expected frequencies
        if (expected < 5).any():
            warning_msg = "Warning: One or more expected frequencies are less than 5. Results may not be reliable."
        else:
            warning_msg = "All expected frequencies are 5 or more."
        
        # Assessing difference between groups' proportions
        reject_null = p_val_chi2 < 0.05
        
        # Create a summary of the test results
        summary_test = (
            f'Chi2 Test for proportions\n'
            f'H0: Equal proportion\n'
            f'chi2: {chi2_stat:.2f}\n'
            f'dof: {dof}\n'
            f'p: {p_val_chi2:.6f}\n'
            f'reject_null: {"True" if reject_null else "False"}\n'
            f'Center Effect: {"Yes" if reject_null else "No"}\n'
        )
        
        # Add newlines at the start for spacing
        num_newlines = 4
        summary_test = '\n' * num_newlines + summary_test

        return (summary_test, p_val_chi2)
    

    except Exception as e:
        num_newlines = 4
        summary_test = '\n' * num_newlines
        summary_test += f"Error: {e}"
        return summary_test

def dependancy_between_bool_transplantaion_center(data, bool_var, doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{bool_var} [2/2]")
    run.bold = True
    run.underline = True
    run.font.size = Pt(18)
    centre_pvalue = []

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    row_cells = table.rows[0].cells
    centre_suivi_categories = sorted(df['V2VX_Centre_Suivi'].unique()) 
    sns.set_theme(style='darkgrid')
    plt.style.use('ggplot')
    width = 0.25
    multiplier = 0
    for i, centre in enumerate(centre_suivi_categories, start=1):
        data = data[data['V2VX_Centre_Suivi'] == centre][bool_var].dropna()
        other_data = [data['V2VX_Centre_Suivi'] != centre][bool_var].dropna()

        # Calculate frequency of each category
        data_counts = data.value_counts().sort_index()
        other_data_counts = other_data.value_counts().sort_index()

        # Ensure both data sets have the same categories
        all_categories = sorted(set(data_counts.index).union(set(other_data_counts.index)))
        data_counts = data_counts.reindex(all_categories, fill_value=0)
        other_data_counts = other_data_counts.reindex(all_categories, fill_value=0)

        # Indices of the categories
        indices = np.arange(len(all_categories))

        plt.figure(figsize=(14, 8))

        # Plotting data
        plt.bar(indices - width/2, data_counts, width, label=f'{centre}')
        plt.bar(indices + width/2, other_data_counts, width, label='Other Centers')

        # Add some text for labels, title, and custom x-axis tick labels, etc.
        plt.ylabel('Counts')
        plt.title(f'Counts by category for {centre} vs Other Centers')
        plt.xticks(indices, all_categories, rotation='vertical')
        plt.legend()

        # Show plot
        plt.tight_layout()
        plt.savefig(f'{save_dir}/{categorical_variable}_per_cat_{centre}.png', dpi=300, bbox_inches='tight')

    for i, centre in enumerate(centre_suivi_categories, start=1):

        # summary_res = f"t_stat: {centre_tvalue[i-1]:.2}\n"
        # summary_res += f"uncorrected_pval: {centre_pvalue[i-1]:.6f}\n"
        # summary_res += f"pval_corrected (Bonferroni): {p_val_corrected:.6f}\n"
        # summary_res += f"dof: {dof}\n"
        # summary_res += f"reject_null: {reject_null}\n"
        # summary_res = '\n'*num_newline + summary_res

      

        # run = row_cells[(i-1)%4].paragraphs[0].add_run(summary_res)
        # run.font.size = Pt(6.3)

        row_cells[(i-1)%4].add_paragraph().add_run().add_picture(f'{save_dir}/{bool_var}_per_cat_{centre}.png', width=Inches(2))


        if i%4 == 0 and i != 12: 
            row_cells = table.add_row().cells
        summary_res = ''
        # if i%2 == 0:
        #     doc.add_paragraph()
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.0)
    doc.add_page_break()

def plot_stacked_bar(df_unique, variable):
    prefixes = ["Admin","V0_Antecedant", "V0_Colonisation_Bact", "V0_Colonisation_Champignon", "V0_Colonisation_MycoBact", "V0_EFR", "V0_GDS", "V0_Generales", "V0_Histocompatibilite", "V0", "SerologiesV0_", "V0_Traitement", "V0_Ventillation"]
   
    data = df_unique.copy()

    grouped_data = data.groupby(['V2VX_Centre_Suivi', variable], observed = False).size()

    unstacked_data = grouped_data.unstack(fill_value=0)

    # Plotting
    sns.set(style='darkgrid', palette='muted')
    plt.style.use('ggplot')

    # Plot as a stacked bar chart
    ax = unstacked_data.plot(kind='bar', stacked=True, figsize=(9, 7), color=['#1f77b4', '#ff7f0e'])
    
    # Remove prefixes from the title
    for prefix in prefixes:
        if prefix in variable:
            cleaned_title = variable.replace(prefix, '').replace('_', ' ')
            continue
    try:
        plt.title(cleaned_title, fontsize=16, fontweight='bold')
    except:
        plt.title(variable, fontsize=16, fontweight='bold')
        print(f'No prefix found in {variable}')
    plt.xlabel('')

    labels = [label.get_text().replace('Colt ', '') for label in ax.get_xticklabels()]
    ax.set_xticklabels(labels, rotation=45, ha='right', fontsize = 12, weight = 'bold')

    # Adding a grid
    plt.grid(axis='y', linestyle='--', alpha=0.7)

    # Adding value labels on top of the bars
    for container in ax.containers:
        ax.bar_label(container, label_type='center')

    
    plt.tight_layout()
    plt.text(1, -0.1, f'n = {data[variable].count()}', 
         horizontalalignment='right', verticalalignment='top', fontsize=12, fontweight='bold',
         transform=plt.gca().transAxes)
    # Move the legend to the top right corner
    plt.legend(loc='upper left', )
    plt.savefig(f'{save_dir}/{variable}_stacked.png')
    plt.close()

def add_centre_comparison_after_CHI2_to_doc(doc, df, variable, centre_suivi_categories):
    # new page
    doc.add_page_break()
    data = df.copy()
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.30)
    section.bottom_margin = Inches(0.30)
    section.left_margin = Inches(0.30)
    section.right_margin = Inches(0.30)

    prefixes_2 = ["Admin", "V0_Antecedant", "V0_Colonisation_Bact", "V0_Colonisation_Champignon", "V0_Colonisation_MycoBact", "V0_EFR", "V0_GDS", "V0_Generales", "V0_Histocompatibilite", "V0_Maladie_Initiale", "SerologiesV0_", "V0_Traitement", "V0_Ventillation"]
    for prefix in prefixes_2:
        if prefix in variable:
            clean_variable = variable.replace(prefix, '').replace('_', ' ')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{variable} [2/2]")
    run.bold = True
    run.underline = True
    run.font.size = Pt(18)
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'

    centre_pvalue = []
    centre_chi2_stats = []
    centre_names = []
    centre_dof = []
    list_variables = []
    
    row_cells = table1.rows[0].cells 
    sns.set_theme(style='darkgrid')
    plt.style.use('ggplot')
    color_palette = sns.color_palette("husl", len(centre_suivi_categories))

       
    for i, centre in enumerate(centre_suivi_categories, start=1):
        center_data = df[df['V2VX_Centre_Suivi'] == centre][variable].dropna()
        other_data = df[df['V2VX_Centre_Suivi'] != centre][variable].dropna()
        try:
            chi2_stat, p_val_chi2, dof, expected = chi2_contingency([center_data.value_counts(), other_data.value_counts()])
            centre_chi2_stats.append(chi2_stat)
            centre_pvalue.append(p_val_chi2)
            centre_dof.append(dof)
            centre_names.append(centre)
        except Exception as e:
            print(f"Variable {variable} Error: {e}")
            list_variables.append(variable)
            continue
        if list_variables:
            print(f"Variables with error: {list_variables}")
        
    
        data_counts = center_data.value_counts().reindex([True, False]).fillna(0)
        other_data_counts = other_data.value_counts().reindex([True, False]).fillna(0)
    
        # Create a DataFrame for plotting
        plot_df = pd.DataFrame({centre.replace('Colt', ''): data_counts, 'Other others': other_data_counts}).T
    
        # Plot
        plot_df.plot(kind='bar', stacked=True, figsize=(10, 6), color=['#ff7f0e','#1f77b4'])
        plt.title(f'{centre.replace("Colt", "")} VS All Others', fontsize=16, fontweight='bold')
        plt.ylabel('Count')
        plt.xticks(rotation=0, fontsize = 18, weight = 'bold')  # Adjust x-ticks to improve readability
        plt.legend(title='Value', labels=['True', 'False'])
        try:
            plt.savefig(f'{save_dir}/{centre}_{clean_variable}_bar.png')
        except:
            plt.savefig(f'{save_dir}/{centre}_{variable}_bar.png')
        plt.close()
        


    reject, pvals_corrected, _, _ = multipletests(centre_pvalue, method='bonferroni')
    for i, (centre, p_val_corrected, reject_null) in enumerate(zip(centre_names, pvals_corrected, reject), start=1):
        summary_res = f"chi2: {centre_chi2_stats[i-1]:.2f}\n"
        summary_res += f"uncorrected_pval: {centre_pvalue[i-1]:.6f}\n"
        summary_res += f"pval_corrected (Bonferroni): {p_val_corrected:.6f}\n"
        summary_res += f"dof: {centre_dof[i-1]}\n"
        summary_res += f"reject_null: {reject_null}\n"
    
        run = row_cells[(i-1)%4].paragraphs[0].add_run(summary_res)
        run.font.size = Pt(6.3)
        try:
            row_cells[(i-1)%4].add_paragraph().add_run().add_picture(f'{save_dir}/{centre}_{clean_variable}_bar.png', width=Inches(2.5))
        except:
            row_cells[(i-1)%4].add_paragraph().add_run().add_picture(f'{save_dir}/{centre}_{variable}_bar.png', width=Inches(2.5))

        if i%4 == 0 and i != 12: 
            row_cells = table1.add_row().cells
        summary_res = ''
        if i%2 == 0:
            doc.add_paragraph()
    for row in table1.rows:
        for cell in row.cells:
            cell.width = Inches(1.0)
    # doc.add_page_break()

def add_stats_bool_to_doc(df, variable, doc):
    data = df.copy()
    # section = doc.sections[-1]
    # new_width, new_height = section.page_height, section.page_width
    # section.orientation = WD_ORIENT.LANDSCAPE
    # section.page_width = new_width
    # section.page_height = new_height
    # section.top_margin = Inches(0.30)
    # section.bottom_margin = Inches(0.30)
    # section.left_margin = Inches(0.30)
    # section.right_margin = Inches(0.30)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{variable} [1/2]")
    run.bold = True
    run.underline = True
    run.font.size = Pt(18)

    centre_suivi_categories = sorted(data['V2VX_Centre_Suivi'].unique())

    all_categories = ['All'] + centre_suivi_categories
    table = doc.add_table(rows=6, cols=len(centre_suivi_categories) + 2)
    table.style = 'Table Grid'

    # Set the headers
    for j, category in enumerate(all_categories, start=1):
        cell = table.cell(0, j)
        cell.text = category.replace('Colt', '') if category != 'All' else category
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(6)

    for i, name in enumerate(['Missing', 'Count', 'True', 'False', 'diag'], start=1):
        cell = table.cell(i, 0)
        cell.text = name
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(5.3)
    table.columns[0].width = Inches(0.53)

    for j, category in enumerate(all_categories, start=1):
        if category == 'All':
            df_category = data
        else:
            df_category = data[data['V2VX_Centre_Suivi'] == category]
        total_count = df_category[variable].count()
        true_count = df_category[variable].sum()
        false_count = total_count - true_count
        true_percentage = (true_count / total_count) * 100 if total_count > 0 else 0
        false_percentage = (false_count / total_count) * 100 if total_count > 0 else 0
        stats = [
            round(df_category[variable].isna().sum(), 2),
            total_count,
            f"{round(true_count, 2)} ({true_percentage:.2f}%)",
            f"{round(false_count, 2)} ({false_percentage:.2f}%)"
        ]
        for i, stat in enumerate(stats, start=1):
            cell = table.cell(i, j)
            cell.text = str(stat)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    
        if category == 'All':
            df_category = data
        if category != 'All':
            df_category = data[data['V2VX_Centre_Suivi'] == category]
        col = variable  
        num_unique_values = df_category[col].nunique()
        num_unique_values = df_category[col].nunique()
        explode = (0.1,) + (0,) * (num_unique_values - 1)
        # colors = sns.color_palette('pastel')[0:num_unique_values]
        def make_autopct(values):
            def my_autopct(pct):
                total = sum(values)
                val = int(round(pct * total / 100.0))
                return '{p:.2f}%  ({v:d})'.format(p=pct, v=val) if pct > 0 else ''
            return my_autopct
        plt.figure(figsize=(6, 6))
        df_category[col].value_counts().plot(
            kind='pie', 
            autopct=make_autopct(df_category[col].value_counts()), 
            color=['#1f77b4', '#ff7f0e'], 
            explode=explode, 
            shadow=True, 
            labels=df_category[col].value_counts().index, 
            textprops=dict(color="black", weight='bold', fontsize=18), 
            labeldistance=1.1
        )
        plt.title(f'{category.replace('Colt', '')}', fontsize=18, weight='bold')
        plt.ylabel('')
        # Add a legend outside the pie chart
        # plt.legend(title=col, bbox_to_anchor=(1.05, 1), loc='best')
        filename = f"{category.replace(' ', '_')}_pie_chart.png"
        plt.savefig(f'{save_dir}/{filename}')
        plt.clf()
        plt.close()
        # Insertion
        try:
            cell = table.cell(5, j)  
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(f'{save_dir}/{filename}', width=Inches(0.8))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception as e:
            print(f"Error: {e}")
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    # Chi2 Independance test
    cell = table.cell(0, 0)
    summary, p_val_chi2 = chi2_summary(data, variable, centre_suivi_categories)
    run = cell.paragraphs[0].add_run(summary)
    run.font.size = Pt(11)
    table.rows[0].cells[0].width = Inches(2.5)
    # reduce the width of the cell
    plot_stacked_bar(data, variable)
    cell = table.cell(0, 1)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(f'{save_dir}/{variable}_stacked.png', width=Inches(6.2), height=Inches(4.2))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    if p_val_chi2 < 0.05:
        # post hoc test
        add_centre_comparison_after_CHI2_to_doc(doc, df, variable, centre_suivi_categories)
    else:
        doc.add_page_break()

#-----------------------------------------------------------------------------------------------------------#

boolean_cols = df_unique.select_dtypes(include='boolean').columns
v0_boolean_cols = [col for col in boolean_cols if col.startswith('V0_')]
v0_boolean_cols = ['V0_Maladie_Initiale_Mucoviscidose', 'V0_Maladie_Initiale_Pneumopathie_Interstitielle_Diffuse']#, 'V0_Maladie_Initiale_BPCO_Emphyseme', 'V0_Maladie_Initiale_HTAP_Primitive', 'V0_Maladie_Initiale_Autres', 'V0_Ventillation_VNI', 'V0_Ventillation_Oxygenotherapie', 'V0_Serologies_Ac_EBV', 'V0_Serologies_Ac_Toxo', 'V0_Serologies_Ac_VHC', 'V0_Serologies_Ag_HBs', 'V0_Serologies_Ac_HBs', 'V0_Serologies_Ac_HBc', 'V0_Serologies_HIV']
for col in v0_boolean_cols:
    add_stats_bool_to_doc(df_unique, col, doc)
# add_stats_bool_to_doc(df_unique, 'V0_Maladie_Initiale_BPCO_Emphyseme', doc)
#------------------------------------------------------------------------------------------------------------#

doc.save("testbool.docx")
os.startfile("testbool.docx")
print(f"--- {time.time() - start_time} seconds ---")